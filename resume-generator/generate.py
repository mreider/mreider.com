#!/usr/bin/env python3
"""
Resume Generator Script
Converts resume text files to styled HTML resumes and PDFs
Supports multiple people: matt, max, alison
"""

import re
import subprocess
import sys
import os
import shutil
import argparse
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Person-specific configurations
PERSON_CONFIGS = {
    'matt': {
        'input_file': 'matt_resume.txt',
        'profile_image': 'matt-profile-image.jpeg',
        'linkedin_url': 'https://linkedin.com/in/mreider',
        'output_dir': 'matt',
        'resume_type': 'professional',
        'show_skills_grid': False,
        'show_certifications': True,
        'show_personal': True,
    },
    'alison': {
        'input_file': 'alison_resume.txt',
        'profile_image': None,
        'linkedin_url': None,
        'output_dir': 'alison',
        'resume_type': 'professional',
        'show_skills_grid': False,
        'show_certifications': False,
        'show_personal': False,
    },
}


def parse_resume_data(filename: str) -> Dict[str, Any]:
    """Parse the resume data from text file"""
    with open(filename, 'r', encoding='utf-8') as f:
        content = f.read()

    data = {}

    # Parse basic info
    data['name'] = re.search(r'NAME: (.+)', content).group(1)
    phone_match = re.search(r'PHONE: (.+)', content)
    data['phone'] = phone_match.group(1).strip() if phone_match and phone_match.group(1).strip() else ""
    email_match = re.search(r'EMAIL: (.+)', content)
    data['email'] = email_match.group(1).strip() if email_match and email_match.group(1).strip() else ""
    data['location'] = re.search(r'LOCATION: (.+)', content).group(1)

    # Parse summary
    summary_match = re.search(r'SUMMARY:\n(.+?)(?=\n\n|\nJOB TITLE|\nEDUCATION|\nLEADERSHIP)', content, re.DOTALL)
    data['summary'] = summary_match.group(1).strip() if summary_match else ""

    # Parse jobs
    data['jobs'] = []

    # Pattern for jobs with bullets
    job_with_bullets_pattern = r'JOB TITLE: (.+?)\nCOMPANY: (.+?)\nDATES: (.+?)\n((?:- .+\n?)+)'
    for match in re.finditer(job_with_bullets_pattern, content):
        title, company, dates, bullets = match.groups()
        bullet_list = [line.strip('- ').strip() for line in bullets.split('\n') if line.strip().startswith('-')]
        data['jobs'].append({
            'title': title,
            'company': company,
            'dates': dates,
            'bullets': bullet_list
        })

    # Pattern for jobs without bullets
    job_without_bullets_pattern = r'JOB TITLE: (.+?)\nCOMPANY: (.*?)\nDATES: (.+?)(?=\n\n|\nJOB TITLE|\nEDUCATION|\nACTIVITIES|\nSKILLS|\nLANGUAGES|$)'
    for match in re.finditer(job_without_bullets_pattern, content):
        title, company, dates = match.groups()
        # Skip if this job was already found with bullets
        already_exists = any(job['title'] == title.strip() and job['company'] == company.strip() for job in data['jobs'])
        if not already_exists:
            data['jobs'].append({
                'title': title.strip(),
                'company': company.strip(),
                'dates': dates.strip(),
                'bullets': []
            })

    # Parse education
    data['education'] = []
    edu_match = re.search(r'EDUCATION:\n\n(.+?)(?=\nLICENSES|\nLEADERSHIP|\nACTIVITIES|\nSKILLS|\nPERSONAL|$)', content, re.DOTALL)
    if edu_match:
        edu_text = edu_match.group(1).strip()
        lines = edu_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line and not line.startswith('Graduated') and not line.startswith('Pupil Personnel') and not line.startswith('High Level') and not line.startswith('An interdisciplinary'):
                degree = line
                school = ""
                date = ""
                details = ""

                # Get school from next line
                if i + 1 < len(lines):
                    school = lines[i + 1].strip()

                # Get graduation date from next line
                if i + 2 < len(lines):
                    next_line = lines[i + 2].strip()
                    if next_line.startswith('Graduated'):
                        date = next_line.replace('Graduated ', '')
                        i += 1
                    elif re.match(r'^[A-Z][a-z]+ \d{4}', next_line):
                        # Date format like "September 2025 - Present"
                        date = next_line
                        i += 1

                # Check for additional details
                if i + 2 < len(lines):
                    detail_line = lines[i + 2].strip()
                    if detail_line and not re.match(r'^[A-Z][a-z]+ of|^Bachelor|^Master|^International', detail_line):
                        details = detail_line
                        i += 1

                data['education'].append({
                    'degree': degree,
                    'school': school,
                    'date': date,
                    'details': details if details else None
                })
                i += 2
            else:
                i += 1

    # Parse leadership training (for Max's resume)
    data['leadership'] = []
    leadership_match = re.search(r'LEADERSHIP TRAINING:\n\n(.+?)(?=\nACTIVITIES|$)', content, re.DOTALL)
    if leadership_match:
        # Use the job parser for leadership section too
        leadership_content = leadership_match.group(1)
        job_pattern = r'JOB TITLE: (.+?)\nCOMPANY: (.+?)\nDATES: (.+?)\n((?:- .+\n?)+)'
        for match in re.finditer(job_pattern, leadership_content):
            title, company, dates, bullets = match.groups()
            bullet_list = [line.strip('- ').strip() for line in bullets.split('\n') if line.strip().startswith('-')]
            data['leadership'].append({
                'title': title,
                'company': company,
                'dates': dates,
                'bullets': bullet_list
            })

    # Parse activities (for Max's resume)
    data['activities'] = []
    activities_match = re.search(r'ACTIVITIES:\n\n(.+?)(?=\nWORK EXPERIENCE|\nSKILLS|$)', content, re.DOTALL)
    if activities_match:
        activities_text = activities_match.group(1).strip()
        # Split by double newlines or activity headers
        blocks = re.split(r'\n\n+', activities_text)
        for block in blocks:
            lines = [l.strip() for l in block.split('\n') if l.strip()]
            if lines:
                activity_name = lines[0]
                activity_bullets = [l.strip('- ').strip() for l in lines[1:] if l.strip().startswith('-')]
                data['activities'].append({
                    'name': activity_name,
                    'bullets': activity_bullets
                })

    # Parse work experience (for student resumes)
    data['work_experience'] = []
    work_match = re.search(r'WORK EXPERIENCE:\n\n(.+?)(?=\nSKILLS|$)', content, re.DOTALL)
    if work_match:
        work_content = work_match.group(1)
        job_pattern = r'JOB TITLE: (.+?)\nCOMPANY: (.+?)\nDATES: (.+?)\n((?:- .+\n?)*)'
        for match in re.finditer(job_pattern, work_content):
            title, company, dates, bullets = match.groups()
            bullet_list = [line.strip('- ').strip() for line in bullets.split('\n') if line.strip().startswith('-')]
            data['work_experience'].append({
                'title': title,
                'company': company,
                'dates': dates,
                'bullets': bullet_list
            })

    # Parse skills (simple list for Max)
    data['skills'] = []
    skills_match = re.search(r'SKILLS:\n((?:- .+\n?)+)', content)
    if skills_match:
        skills_text = skills_match.group(1)
        data['skills'] = [line.strip('- ').strip() for line in skills_text.split('\n') if line.strip().startswith('-')]

    # Parse languages
    languages_match = re.search(r'LANGUAGES:\n(.+?)(?=\n\n|$)', content, re.DOTALL)
    data['languages'] = languages_match.group(1).strip() if languages_match else ""

    # Parse interests
    interests_match = re.search(r'INTERESTS:\n(.+?)(?=\n\n|$)', content, re.DOTALL)
    data['interests'] = interests_match.group(1).strip() if interests_match else ""

    # Parse references
    references_match = re.search(r'REFERENCES:\n(.+?)(?=\n\n|$)', content, re.DOTALL)
    data['references'] = references_match.group(1).strip() if references_match else ""

    # Parse competencies (for Matt's resume)
    comp_match = re.search(r'CORE COMPETENCIES:\n(.+?)(?=\nEDUCATION)', content, re.DOTALL)
    if comp_match:
        comp_text = comp_match.group(1).strip()
        data['competencies'] = [line.strip('• ').strip() for line in comp_text.split('\n') if line.strip().startswith('•')]
    else:
        data['competencies'] = []

    # Parse licenses & certifications
    data['certifications'] = []
    cert_match = re.search(r'LICENSES & CERTIFICATIONS:\n\n(.+?)(?=\nPERSONAL|$)', content, re.DOTALL)
    if cert_match:
        cert_text = cert_match.group(1).strip()
        lines = cert_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line and not line.startswith('Issued'):
                cert_name = line
                organization = ""
                date = ""

                if i + 1 < len(lines):
                    organization = lines[i + 1].strip()

                if i + 2 < len(lines) and lines[i + 2].strip().startswith('Issued'):
                    date = lines[i + 2].strip().replace('Issued ', '')

                data['certifications'].append({
                    'name': cert_name,
                    'organization': organization,
                    'date': date
                })
                i += 3
            else:
                i += 1

    # Parse personal
    personal_match = re.search(r'PERSONAL:\n(.+)', content, re.DOTALL)
    data['personal'] = personal_match.group(1).strip() if personal_match else ""

    return data


def generate_html(data: Dict[str, Any], config: Dict[str, Any]) -> str:
    """Generate simple ATS-friendly HTML resume - single column, no fancy formatting"""

    # Build jobs HTML
    jobs_html = ""
    for job in data['jobs']:
        bullets_html = ""
        for bullet in job.get('bullets', []):
            bullets_html += f"<li>{bullet}</li>\n"

        company_line = f" - {job['company']}" if job['company'] else ""
        jobs_html += f"""
        <div class="job">
            <div class="job-header">
                <strong>{job['title']}{company_line}</strong>
                <span class="dates">{job['dates']}</span>
            </div>
            <ul>
{bullets_html}            </ul>
        </div>
"""
        # Add page break after Macromedia
        if 'Macromedia' in job.get('company', ''):
            jobs_html += '        <div class="page-break"></div>\n'

    # Build education HTML
    edu_html = ""
    for edu in data['education']:
        edu_html += f"""
        <div class="education">
            <strong>{edu['degree']}</strong> - {edu['school']}, {edu['date']}
        </div>
"""

    # Build certifications HTML
    cert_html = ""
    if config.get('show_certifications') and data.get('certifications'):
        for cert in data['certifications']:
            cert_html += f"""
        <div class="cert">
            <strong>{cert['name']}</strong> - {cert['organization']}, {cert['date']}
        </div>
"""

    # Build skills HTML
    skills_html = ", ".join(data.get('skills', []))

    # Build contact lines
    contact_lines = []
    if data.get('email'):
        contact_lines.append(data['email'])
    if data.get('phone'):
        # Split multiple phone numbers on |
        phones = [p.strip() for p in data['phone'].split('|')]
        contact_lines.extend(phones)
    if data.get('location'):
        contact_lines.append(data['location'])
    contact_html = "<br>\n        ".join(contact_lines)

    # Simple ATS-friendly HTML template
    html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{data['name']} - Resume</title>
    <style>
        body {{
            font-family: Arial, Helvetica, sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #000000;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 0.5in;
            background: #ffffff;
        }}
        h1 {{
            font-size: 18pt;
            margin: 0 0 5px 0;
            border-bottom: 1px solid #000;
            padding-bottom: 5px;
        }}
        .contact {{
            margin-bottom: 15px;
            font-size: 10pt;
        }}
        h2 {{
            font-size: 12pt;
            margin: 18px 0 10px 0;
            text-transform: uppercase;
            border-bottom: 1px solid #666;
            padding-bottom: 3px;
        }}
        .summary {{
            margin-bottom: 18px;
        }}
        .job {{
            margin-bottom: 16px;
            page-break-inside: avoid;
        }}
        .job-header {{
            display: flex;
            justify-content: space-between;
            margin-bottom: 3px;
        }}
        .dates {{
            font-style: italic;
            color: #333;
        }}
        ul {{
            margin: 5px 0 0 20px;
            padding: 0;
        }}
        li {{
            margin-bottom: 4px;
        }}
        .education, .cert {{
            margin-bottom: 8px;
        }}
        .skills {{
            margin-top: 8px;
        }}
        .page-break {{
            page-break-before: always;
        }}
        @media print {{
            body {{ padding: 0.25in; }}
            .page-break {{ page-break-before: always; }}
        }}
    </style>
</head>
<body>
    <h1>{data['name']}</h1>
    <div class="contact">
        {contact_html}
    </div>

    <h2>Summary</h2>
    <div class="summary">{data['summary']}</div>

    <h2>Experience</h2>
{jobs_html}

    <h2>Education</h2>
{edu_html}

    {"<h2>Certifications</h2>" + cert_html if cert_html else ""}

    {"<h2>Skills</h2><div class='skills'>" + skills_html + "</div>" if skills_html else ""}
</body>
</html>"""

    return html_template


def generate_pdf(output_dir: str):
    """Generate PDF from HTML using Chrome headless"""
    html_path = os.path.join(output_dir, 'index.html')
    pdf_path = os.path.join(output_dir, 'resume.pdf')

    methods = [
        ['/Applications/Google Chrome.app/Contents/MacOS/Google Chrome', '--headless', '--disable-gpu', f'--print-to-pdf={pdf_path}', '--no-margins', '--no-pdf-header-footer', html_path],
        ['google-chrome', '--headless', '--disable-gpu', f'--print-to-pdf={pdf_path}', '--no-margins', '--no-pdf-header-footer', html_path],
        ['chromium-browser', '--headless', '--disable-gpu', f'--print-to-pdf={pdf_path}', '--no-margins', '--no-pdf-header-footer', html_path],
    ]

    for method in methods:
        try:
            result = subprocess.run(method, capture_output=True, text=True)
            if result.returncode == 0:
                print(f"  PDF generated using {method[0]}")
                return True
        except FileNotFoundError:
            continue

    print(f"  Could not generate PDF automatically for {output_dir}")
    return False


def generate_docx(data: Dict[str, Any], config: Dict[str, Any], output_dir: str):
    """Generate Word document from resume data"""
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Name header
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(data['name'])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Contact info
    contact_para = doc.add_paragraph()
    contact_lines = []
    if data.get('email'):
        contact_lines.append(data['email'])
    if data.get('phone'):
        phones = [p.strip() for p in data['phone'].split('|')]
        contact_lines.extend(phones)
    if data.get('location'):
        contact_lines.append(data['location'])
    contact_para.add_run("\n".join(contact_lines))
    contact_para.paragraph_format.space_after = Pt(12)

    # Summary section
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run("SUMMARY")
    summary_run.bold = True
    summary_run.font.size = Pt(12)
    summary_para = doc.add_paragraph(data['summary'])
    summary_para.paragraph_format.space_after = Pt(12)

    # Experience section
    exp_heading = doc.add_paragraph()
    exp_run = exp_heading.add_run("EXPERIENCE")
    exp_run.bold = True
    exp_run.font.size = Pt(12)

    for job in data['jobs']:
        job_para = doc.add_paragraph()
        company_line = f" - {job['company']}" if job['company'] else ""
        title_run = job_para.add_run(f"{job['title']}{company_line}")
        title_run.bold = True
        job_para.add_run(f"  |  {job['dates']}")
        job_para.paragraph_format.space_after = Pt(2)

        for bullet in job.get('bullets', []):
            bullet_para = doc.add_paragraph(style='List Bullet')
            bullet_para.add_run(bullet)
            bullet_para.paragraph_format.space_after = Pt(2)

    # Education section
    edu_heading = doc.add_paragraph()
    edu_run = edu_heading.add_run("EDUCATION")
    edu_run.bold = True
    edu_run.font.size = Pt(12)

    for edu in data['education']:
        edu_para = doc.add_paragraph()
        edu_title = edu_para.add_run(f"{edu['degree']}")
        edu_title.bold = True
        edu_para.add_run(f" - {edu['school']}, {edu['date']}")
        edu_para.paragraph_format.space_after = Pt(4)

    # Certifications section
    if config.get('show_certifications') and data.get('certifications'):
        cert_heading = doc.add_paragraph()
        cert_run = cert_heading.add_run("CERTIFICATIONS")
        cert_run.bold = True
        cert_run.font.size = Pt(12)

        for cert in data['certifications']:
            cert_para = doc.add_paragraph()
            cert_title = cert_para.add_run(f"{cert['name']}")
            cert_title.bold = True
            cert_para.add_run(f" - {cert['organization']}, {cert['date']}")
            cert_para.paragraph_format.space_after = Pt(4)

    # Skills section
    if data.get('skills'):
        skills_heading = doc.add_paragraph()
        skills_run = skills_heading.add_run("SKILLS")
        skills_run.bold = True
        skills_run.font.size = Pt(12)
        skills_para = doc.add_paragraph(", ".join(data['skills']))

    # Save the document
    docx_path = os.path.join(output_dir, 'resume.docx')
    doc.save(docx_path)
    print(f"  Word document written to: {docx_path}")
    return True


def generate_index_page(script_dir: str, hugo_static_resume: str):
    """Generate the main index page for /resume/"""
    index_html = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Matthew Reider - Resume</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Crimson+Text:wght@400;600&display=swap');

        body {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif;
            line-height: 1.6;
            color: #000000;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            background: #ffffff;
        }

        h1 {
            font-family: 'Crimson Text', serif;
            font-size: 28pt;
            font-weight: 700;
            margin-bottom: 30px;
            padding-bottom: 15px;
            border-bottom: 2px solid #000000;
            letter-spacing: -0.5pt;
        }

        .resume-list {
            list-style: none;
            padding: 0;
            margin: 0;
        }

        .resume-item {
            margin-bottom: 20px;
            padding: 20px 25px;
            background: #f7f7f7;
            border: 1px solid #cccccc;
            border-left: 4px solid #000000;
            transition: all 0.2s ease;
        }

        .resume-item:hover {
            background: #f0f0f0;
            transform: translateX(5px);
        }

        .resume-item a {
            text-decoration: none;
            color: inherit;
            display: block;
        }

        .resume-name {
            font-family: 'Crimson Text', serif;
            font-size: 16pt;
            font-weight: 600;
            color: #000000;
            margin-bottom: 5px;
        }

        .resume-title {
            font-size: 11pt;
            color: #555555;
        }

        .resume-links {
            margin-top: 10px;
            display: flex;
            gap: 15px;
        }

        .resume-links a {
            font-size: 10pt;
            color: #0066cc;
            text-decoration: none;
        }

        .resume-links a:hover {
            text-decoration: underline;
        }
    </style>
</head>
<body>
    <h1>Resume</h1>
    <ul class="resume-list">
        <li class="resume-item">
            <a href="matt/">
                <div class="resume-name">Matthew Reider</div>
                <div class="resume-title">Principal Product Manager</div>
            </a>
            <div class="resume-links">
                <a href="matt/">View Resume</a>
                <a href="matt/resume.pdf">Download PDF</a>
            </div>
        </li>
    </ul>
</body>
</html>"""

    index_path = os.path.join(hugo_static_resume, 'index.html')
    with open(index_path, 'w', encoding='utf-8') as f:
        f.write(index_html)
    print(f"Index page created at: {index_path}")


def generate_resume(person: str, script_dir: str, hugo_static_resume: str):
    """Generate resume for a specific person"""
    config = PERSON_CONFIGS[person]
    input_file = os.path.join(script_dir, config['input_file'])
    output_subdir = os.path.join(hugo_static_resume, config['output_dir'])

    print(f"\n{'='*50}")
    print(f"Generating resume for: {person.upper()}")
    print(f"{'='*50}")

    # Create output directory
    os.makedirs(output_subdir, exist_ok=True)

    # Parse resume data
    print(f"  Reading {config['input_file']}...")
    data = parse_resume_data(input_file)

    # Generate HTML
    print("  Generating HTML...")
    html = generate_html(data, config)

    # Write index.html
    html_path = os.path.join(output_subdir, 'index.html')
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f"  HTML written to: {html_path}")

    # Generate PDF
    print("  Generating PDF...")
    generate_pdf(output_subdir)

    # Generate Word document
    print("  Generating Word document...")
    generate_docx(data, config, output_subdir)

    print(f"  Resume generated for {person}")


def main():
    """Main function to generate all resumes"""
    parser = argparse.ArgumentParser(description='Generate resumes for the Reider family')
    parser.add_argument('--person', choices=['matt', 'alison', 'all'], default='all',
                        help='Which person to generate resume for (default: all)')
    args = parser.parse_args()

    script_dir = os.path.dirname(os.path.abspath(__file__))
    hugo_static_resume = os.path.join(script_dir, '..', 'static', 'resume')

    # Create main resume directory
    os.makedirs(hugo_static_resume, exist_ok=True)

    try:
        if args.person == 'all':
            # Generate all resumes
            for person in PERSON_CONFIGS.keys():
                generate_resume(person, script_dir, hugo_static_resume)

        else:
            # Generate single resume
            generate_resume(args.person, script_dir, hugo_static_resume)

        print(f"\n{'='*50}")
        print("Resume generation complete!")
        print(f"Output directory: {hugo_static_resume}")
        print(f"{'='*50}")

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0


if __name__ == "__main__":
    exit(main())
